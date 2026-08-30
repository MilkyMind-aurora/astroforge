"""MD → DOCX 转换器（python-docx 实现骨架：覆盖方案 3.4.2 核心元素）。

完整版（自动目录/页眉页脚/图片编号/中文排版优化）属 Phase 4.2.4。
"""
from __future__ import annotations

import re
from pathlib import Path

from docx.shared import Pt, RGBColor
from template_manager import load_template

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def _apply_inline(paragraph, text: str) -> None:
    """行内格式：**粗** *斜* `行内码`。"""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("*") and token.endswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0x88, 0x00, 0x00)
        else:
            paragraph.add_run(token)


def _add_code_block(document, lines: list[str]) -> None:
    """代码块：等宽字体 + 浅灰底（段落底纹用字符底色近似，保留缩进）。"""
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Pt(12)
        run = paragraph.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def _add_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, cell in enumerate(row):
            if c < len(table.rows[r].cells):
                table.rows[r].cells[c].text = cell.strip()


def convert_lines(document, lines: list[str]) -> None:
    """逐行解析 Markdown 元素写入 document（标题/段落/列表/代码/引用/表格）。"""
    in_code = False
    code_buffer: list[str] = []
    table_buffer: list[list[str]] = []

    def flush_table() -> None:
        if table_buffer:
            _add_table(document, table_buffer)
            table_buffer.clear()

    for line in lines:
        stripped = line.rstrip()
        if stripped.strip().startswith("```"):
            if in_code:
                _add_code_block(document, code_buffer)
                code_buffer.clear()
            in_code = not in_code
            flush_table()
            continue
        if in_code:
            code_buffer.append(line)
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c for c in stripped.strip("|").split("|")]
            if not set(stripped.replace("|", "").strip()) <= {"-", ":", " "}:
                table_buffer.append(cells)
            continue
        flush_table()
        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            document.add_heading(heading.group(2), level=len(heading.group(1)))
        elif stripped.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(24)
            run = paragraph.add_run(stripped[2:])
            run.italic = True
        elif re.match(r"^[-*+]\s+", stripped):
            _apply_inline(document.add_paragraph(style="List Bullet"), re.sub(r"^[-*+]\s+", "", stripped))
        elif re.match(r"^\d+\.\s+", stripped):
            _apply_inline(document.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s+", "", stripped))
        elif stripped.strip():
            _apply_inline(document.add_paragraph(), stripped.strip())
    flush_table()
    if code_buffer:  # 未闭合代码块兜底
        _add_code_block(document, code_buffer)


def convert_file(md_path: Path, docx_path: Path, template: str = "tech_doc") -> Path:
    document = load_template(template)
    convert_lines(document, md_path.read_text(encoding="utf-8").splitlines())
    document.save(docx_path)
    return docx_path


def convert_merged(md_paths: list[Path], docx_path: Path, template: str = "tech_doc") -> Path:
    document = load_template(template)
    for i, md in enumerate(md_paths):
        if i > 0:
            document.add_page_break()
        document.add_heading(md.stem, level=1)
        convert_lines(document, md.read_text(encoding="utf-8").splitlines())
    document.save(docx_path)
    return docx_path
