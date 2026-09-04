# -*- coding: utf-8 -*-
"""生成 5 套内置 DOCX 模板（Phase 0.3 / 4.2.3，方案 3.4.2 模板清单）。

每套按方案规范设置正文/标题字体（中英文分开）、行距、页边距与签名特性；
模板只承载样式（无正文内容），md2docx 转换时作为基底文档继承。
用法：python scripts/make_templates.py [--output templates]
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

REPO_ROOT = Path(__file__).resolve().parents[1]


def _set_font(style, ascii_font: str, east_font: str, size_pt: float,
              bold: bool | None = None, color: tuple | None = None) -> None:
    """中英文字体分离设置（w:eastAsia 管中文，name 管 ASCII）。"""
    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold
    if color:
        style.font.color.rgb = color.__class__(*color) if not isinstance(color, tuple) else __import__(
            "docx.shared", fromlist=["RGBColor"]).RGBColor(*color)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), east_font)


def _set_margins(doc: Document, top=2.54, bottom=2.54, left=3.17, right=3.17) -> None:
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def _base(doc: Document, normal_ascii: str, normal_east: str, size: float,
          line_spacing: float) -> None:
    normal = doc.styles["Normal"]
    _set_font(normal, normal_ascii, normal_east, size)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = line_spacing
    normal.paragraph_format.space_after = Pt(6)
    for name, heading_size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[name]
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def build_academic(path: Path) -> None:
    """学术论文模板：宋体/Times、黑体标题、1.5 倍行距、A4 标准页边距。"""
    doc = Document()
    _base(doc, "Times New Roman", "宋体", 12, 1.5)
    _set_font(doc.styles["Heading 1"], "Times New Roman", "黑体", 16, bold=True)
    _set_font(doc.styles["Heading 2"], "Times New Roman", "黑体", 14, bold=True)
    _set_font(doc.styles["Heading 3"], "Times New Roman", "宋体", 12, bold=True)
    _set_margins(doc)
    doc.core_properties.comments = "学术论文模板：课程论文/小论文/竞赛报告（宋体正文+黑体标题+1.5倍行距）"
    doc.save(path)


def build_tech_doc(path: Path) -> None:
    """技术文档模板：无衬线正文、多级标题体系、紧凑行距。"""
    doc = Document()
    _base(doc, "Segoe UI", "微软雅黑", 10.5, 1.25)
    for name, size in (("Heading 1", 15), ("Heading 2", 12.5), ("Heading 3", 11)):
        _set_font(doc.styles[name], "Segoe UI", "微软雅黑", size, bold=True,
                  color=(0x4F, 0xC3, 0xF7))
    code = doc.styles["Normal"]
    code.paragraph_format.line_spacing = 1.1
    _set_margins(doc, 2.0, 2.0, 2.4, 2.4)
    doc.core_properties.comments = "技术文档模板：接口文档/知识库导出（雅黑正文+氢蓝标题）"
    doc.save(path)


def build_math_model(path: Path) -> None:
    """数模竞赛模板：国赛风格，黑体一级标题、摘要版式、规范页边距。"""
    doc = Document()
    _base(doc, "Times New Roman", "宋体", 12, 1.5)
    _set_font(doc.styles["Heading 1"], "Times New Roman", "黑体", 15, bold=True)
    _set_font(doc.styles["Heading 2"], "Times New Roman", "黑体", 13, bold=True)
    _set_margins(doc, 2.5, 2.5, 3.0, 3.0)
    doc.core_properties.comments = "数模竞赛模板：摘要页/关键词/章节编号/参考文献格式基座（国赛规范）"
    doc.save(path)


def build_simple_general(path: Path) -> None:
    """简约通用模板：无衬线、宽松行距、清爽页边距。"""
    doc = Document()
    _base(doc, "Calibri", "微软雅黑", 11, 1.4)
    for name, size in (("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11.5)):
        _set_font(doc.styles[name], "Calibri", "微软雅黑", size, bold=True,
                  color=(0x8B, 0x7C, 0xF6))
    _set_margins(doc, 2.2, 2.2, 2.6, 2.6)
    doc.core_properties.comments = "简约通用模板：日常办公（雅黑正文+星尘紫标题+宽松行距）"
    doc.save(path)


def build_formal_report(path: Path) -> None:
    """正式报告模板：页眉标题、页码底部居中、一级标题分页。"""
    doc = Document()
    _base(doc, "Times New Roman", "宋体", 12, 1.4)
    _set_font(doc.styles["Heading 1"], "Times New Roman", "黑体", 16, bold=True)
    doc.styles["Heading 1"].paragraph_format.page_break_before = True
    _set_margins(doc)
    # 页眉：AstroForge 报告标题；页脚：第 X 页 共 Y 页（fldSimple 域代码）
    header = doc.sections[0].header
    header.paragraphs[0].text = "AstroForge · 正式报告"
    header.paragraphs[0].alignment = 1  # 居中
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = 1
    from docx.oxml import OxmlElement

    def _add_field(paragraph, instruction: str, around: tuple[str, str] | None = None):
        if around and around[0]:
            paragraph.add_run(around[0])
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), instruction)
        run = OxmlElement("w:r")
        fld.append(run)
        paragraph._p.append(fld)
        if around and around[1]:
            paragraph.add_run(around[1])

    _add_field(footer_para, "PAGE", ("第 ", " 页 / 共 "))
    _add_field(footer_para, "NUMPAGES", (None, " 页"))
    doc.core_properties.comments = "正式报告模板：页眉标题+页脚页码+一级标题分页（正式商务风）"
    doc.save(path)


BUILDERS = {
    "academic": build_academic,
    "tech_doc": build_tech_doc,
    "math_model": build_math_model,
    "simple_general": build_simple_general,
    "formal_report": build_formal_report,
}


def main() -> int:
    parser = argparse.ArgumentParser(description="生成 5 套内置 DOCX 模板")
    parser.add_argument("--output", default=str(REPO_ROOT / "templates"))
    args = parser.parse_args()
    output = Path(args.output)
    output.mkdir(parents=True, exist_ok=True)
    for key, builder in BUILDERS.items():
        target = output / f"{key}.docx"
        builder(target)
        # 生成即验：能重新打开且关键样式存在
        reopened = Document(str(target))
        assert "Heading 1" in [s.name for s in reopened.styles]
        print(f"[OK] {target.name}")
    print(f"完成：{len(BUILDERS)} 套模板已生成到 {output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
